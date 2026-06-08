from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Sayfa kenar boşlukları ─────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Stil yardımcıları ──────────────────────────────────────────────────────

def set_font(run, name='Times New Roman', size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def para(text='', style='Normal', align=WD_ALIGN_PARAGRAPH.LEFT,
         size=12, bold=False, italic=False, space_before=0, space_after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic)
    return p

def heading1(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(f'{num}.\t{text}')
    set_font(run, size=12, bold=True)
    return p

def heading2(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.5)
    run = p.add_run(f'{num}\t{text}')
    set_font(run, size=12, bold=True)
    return p

def heading3(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(1.0)
    run = p.add_run(f'{num}\t{text}')
    set_font(run, size=12, bold=False, italic=False)
    run.font.bold = True
    return p

def body(text, indent=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.left_indent  = Cm(indent)
    run = p.add_run(text)
    set_font(run, size=12)
    return p

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # Header satırı
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=11, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Header arka planı gri
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9D9D9')
        tcPr.append(shd)
    # Veri satırları
    for ri, row_data in enumerate(rows):
        row = t.rows[ri+1]
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(cell_text))
            set_font(run, size=11)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Sütun genişlikleri
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

def page_break():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# KAPAK SAYFASI
# ══════════════════════════════════════════════════════════════════════════

# Üst boşluk
for _ in range(3):
    para()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run('YBS450 Mesleki Eğitim Stajı')
set_font(run, size=16, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(30)
run = p.add_run('Staj Sonuç Raporu')
set_font(run, size=16, bold=True)

for _ in range(2):
    para()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run('Hazırlayan')
set_font(run, size=13, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(40)
run = p.add_run('22298002 Durmuş Emre Uludağ')
set_font(run, size=13)

for _ in range(4):
    para()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run('16 Şubat 2026 – 5 Haziran 2026')
set_font(run, size=12)

page_break()

# ══════════════════════════════════════════════════════════════════════════
# İÇİNDEKİLER
# ══════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
run = p.add_run('İÇİNDEKİLER')
set_font(run, size=14, bold=True)

toc_items = [
    ('İÇİNDEKİLER', '2'),
    ('1.\tİŞBAŞI EĞİTİMİNİN YAPILDIĞI KURUM HAKKINDA BİLGİ', '3'),
    ('\t1.1.\tKURUM HAKKINDA BİLGİ', '3'),
    ('\t1.2.\tKURUMUN FAALİYET ALANLARI', '3'),
    ('\t\t1.2.1\tBilişim Teknolojileri İle İlgili Faaliyet Alanları, Hizmetler ve Projeler', '3'),
    ('\t\t1.2.2\tHizmet Verdiği Sektörler ve Faaliyet Alanları', '3'),
    ('\t1.3.\tKURUMUN ORGANİZASYON YAPISI', '3'),
    ('\t1.4.\tSTAJYERİN BİRLİKTE ÇALIŞTIĞI EKİP VE PROJEDE GÖREV ALANLAR', '3'),
    ('2.\tİŞLETMEDE ÜSTLENİLEN GÖREVLER, YÜRÜTÜLEN FAALİYETLER VE PROJELER', '4'),
    ('\t2.1.\tİDARİ GÖREVLER', '4'),
    ('\t2.2.\tTEKNİK GÖREVLER', '4'),
    ('\t2.3.\tPROJE GÖREVLERİ', '4'),
    ('\t2.4.\tETKİNLİK VE İŞ AKIŞ DİYAGRAMLARI', '4'),
    ('\t\t2.4.1\tGünlük Faaliyetlere Yönelik Etkinlik ve İş Akış Diyagramları', '5'),
    ('\t\t2.4.2\tProje Faaliyetlerine Yönelik Etkinlik ve İş Akış Diyagramları', '5'),
    ('3.\tBÖLÜMDE EĞİTİM SÜRESİNCE ALINAN DERSLERDEN STAJDA YARARLANMA DURUMU', '7'),
    ('4.\tİŞLETMEDEKİ OLUMLU VE/VEYA OLUMSUZ KONULAR HAKKINDA GENEL YORUM', '8'),
    ('\t4.1.\tOLUMLU KONULAR', '8'),
    ('\t4.2.\tOLUMSUZ KONULAR', '8'),
    ('5.\tÖNERİLER', '8'),
    ('\t5.1.\tİŞLETMEDEKİ SORUNLARIN GİDERİLMESİNE YÖNELİK ÖNERİLER', '8'),
    ('\t5.2.\tSTAJ EĞİTİM UYGULAMASINA YÖNELİK ÖNERİLER', '8'),
    ('\t5.3.\tDİĞER ÖNERİLER', '8'),
    ('6.\tSTAJ YOKLAMASI BEYANI', '9'),
]

for label, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(label)
    set_font(run, size=11)
    # Tab + sayfa numarası sağa
    run2 = p.add_run(f'\t{page}')
    set_font(run2, size=11)
    # Sağ hizalamalı tab dur
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:leader'), 'dot')
    tab.set(qn('w:pos'), '8640')  # ~15.24cm
    tabs.append(tab)
    pPr.append(tabs)

page_break()

# ══════════════════════════════════════════════════════════════════════════
# BÖLÜM 1
# ══════════════════════════════════════════════════════════════════════════

heading1('1', 'İşbaşı Eğitiminin Yapıldığı Kurum Hakkında Bilgi')

heading2('1.1.', 'Kurum Hakkında Bilgi')
body('Kurumun adı, yönetim ve kurucuları, organizasyon yapısı, yazışma adresi, telefon, e-posta '
     've yerleşkeleri hakkında bilgi verilecektir.')
para()

heading2('1.2.', 'Kurumun Faaliyet Alanları')
para()

heading3('1.2.1', 'Bilişim Teknolojileri İle İlgili Faaliyet Alanları, Hizmetler ve Projeler')
body('Kurumun faaliyetleri ve projeleri bilişim teknolojileri açısından ele alınacaktır.')
para()

heading3('1.2.2', 'Hizmet Verdiği Sektörler ve Faaliyet Alanları')
body('Kurumun doğrudan ve dolaylı olarak hizmet verdiği sektörlerden bahsedilecektir.')
para()

heading2('1.3.', 'Kurumun Organizasyon Yapısı')
body('Kurumu organizasyon yapısı ile gerçekleştirdiği işletme ve yönetim fonksiyonlarından bahsedilecektir.')
para()

heading2('1.4.', 'Stajyerin Birlikte Çalıştığı Ekip ve Projede Görev Alanlar')
body('Aşağıdaki tabloda projede yer alan takım üyelerine ait bilgiler, görev ve üstlendikleri sorumluluklar belirtilecektir.')
para()

add_table(
    headers=['S. Nu.', 'Görevi', 'Adı ve Soyadı'],
    rows=[
        ['1', 'Takım Lideri', ''],
        ['2', 'Takım Üyesi', ''],
        ['3', 'Takım Üyesi', ''],
        ['4', 'Takım Üyesi', ''],
        ['5', 'Takım Üyesi', ''],
    ],
    col_widths=[1.5, 5.0, 7.5]
)

# ══════════════════════════════════════════════════════════════════════════
# BÖLÜM 2
# ══════════════════════════════════════════════════════════════════════════

heading1('2', 'İşletmede Üstlenilen Görevler, Yürütülen Faaliyetler ve Projeler')
body('Kurumdaki görev ve pozisyona göre aşağıdaki başlıklar hakkında bilgi verilecektir.')
para()

heading2('2.1.', 'İdari Görevler')
body('Üstlenilen idari görevlerden bahsedilecektir.')
para()
para()

heading2('2.2.', 'Teknik Görevler')
body('Bilişim teknolojileri kapsamında işbaşı eğitimi süresince gerçekleştirdiği teknik görev ve faaliyetlerden '
     'bahsedilecektir (yazılım, sistem, donamım vb.).')
para()
para()

heading2('2.3.', 'Proje Görevleri')
body('Proje yönetimi kapsamında işbaşı eğitimi süresince yer aldığı projeler hakkında genel bilgi verilecek '
     '(niteliği, amacı, kapsamı, süresi vb.) ve projelerde üstlenilen görevlerden bahsedecektir.')
para()
para()

heading2('2.4.', 'Etkinlik ve İş Akış Diyagramları')
body('Kurumdaki üstlenilen görevler dikkate alınarak aşağıdaki alt başlıklarda etkinlik (activity) ve '
     'iş akışı (business flow) diyagramları kullanılacaktır.')
para()

heading3('2.4.1', 'Günlük Faaliyetlere Yönelik Etkinlik ve İş Akış Diyagramları')
para()
# Diyagram için boş alan
for _ in range(8):
    para()

page_break()

heading3('2.4.2', 'Proje Faaliyetlerine Yönelik Etkinlik ve İş Akış Diyagramları')
para()
# Diyagram için boş alan
for _ in range(12):
    para()

page_break()

# ══════════════════════════════════════════════════════════════════════════
# BÖLÜM 3
# ══════════════════════════════════════════════════════════════════════════

heading1('3', 'Bölümde Eğitim Süresince Alınan Derslerden Stajda Yararlanma Durumu')
body('Bölümde alınan ders ve bilgilerden işbaşında eğitimde yararlanma durumunu belirtiniz '
     '(hangi konularda hangi derslerden yararlanıldı vb).')
para()

add_table(
    headers=['Dersin Kodu', 'Dersin Adı', 'Dersin Konuları', 'İş Başı Staj Eğitimindeki Görevlere Katkısı'],
    rows=[['', '', '', ''] for _ in range(8)],
    col_widths=[2.5, 3.5, 4.5, 4.0]
)

page_break()

# ══════════════════════════════════════════════════════════════════════════
# BÖLÜM 4
# ══════════════════════════════════════════════════════════════════════════

heading1('4', 'İşletmedeki Olumlu ve/veya Olumsuz Konular Hakkında Genel Yorum')
para()

heading2('4.1.', 'Olumlu Konular')
body('Staj esnasında gözlemlediğiniz olumlu konuları listeleyiniz, varsa sizin mesleki ve kişisel '
     'gelişiminize olan katkıları belirtiniz.')
para()
para()
para()

heading2('4.2.', 'Olumsuz Konular')
body('Staj esnasında gözlemlediğiniz olumsuz konuları listeleyiniz, varsa sizin mesleki ve kişisel '
     'gelişiminize olan etkilerini belirtiniz.')
para()
para()
para()

# ══════════════════════════════════════════════════════════════════════════
# BÖLÜM 5
# ══════════════════════════════════════════════════════════════════════════

heading1('5', 'Öneriler')
para()

heading2('5.1.', 'İşletmedeki Sorunların Giderilmesine Yönelik Öneriler')
body('İşletmede karşılaştığınız olumsuz konular ve problemlere yönelik önerilerinizi belirtiniz.')
para()
para()

heading2('5.2.', 'Staj Eğitim Uygulamasına Yönelik Öneriler')
body('Staj uygulamasıyla ilgili olumlu ve/veya olumsuz konulara yönelik önerilerinizi belirtiniz.')
para()
para()

heading2('5.3.', 'Diğer Öneriler')
body('Varsa diğer konularla ilgili önerilerinizi belirtiniz.')
para()
para()

# ══════════════════════════════════════════════════════════════════════════
# BÖLÜM 6
# ══════════════════════════════════════════════════════════════════════════

heading1('6', 'Staj Yoklaması Beyanı')
body('Bu bölüme stajın ilk sekiz haftasına yönelik telefon ve haftalık günlük yoklama durumu ile '
     'yoklamada bulunulmayan ve/veya işaretlenmeyen günlerle ilgili gerekçeler belirtilecektir '
     '(görev, unutma, sağlık vb.)')
para()

add_table(
    headers=['Hafta Numarası', 'Tarih', 'Katılma Durumu (Var/Yok)', 'Gerekçe'],
    rows=[[str(i), '', '', ''] for i in range(1, 16)],
    col_widths=[3.0, 3.5, 4.5, 4.0]
)

# ── Kaydet ────────────────────────────────────────────────────────────────
output = r'C:\Users\EMRE\Desktop\Staj_Raporu_Emre_Uludag.docx'
doc.save(output)
print(f'Kaydedildi: {output}')
